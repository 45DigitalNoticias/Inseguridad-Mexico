# -*- coding: utf-8 -*-
"""
Genera el Word del DIRECTORIO DE DIFUSIÓN del dashboard "Inseguridad México".
Salida: Directorio de difusion - Inseguridad Mexico.docx (en esta misma carpeta).
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- paleta ----------
NAVY   = RGBColor(0x0D, 0x11, 0x17)
AMBER  = RGBColor(0xB4, 0x7A, 0x10)
GREY   = RGBColor(0x55, 0x5B, 0x63)
INK    = RGBColor(0x1A, 0x1E, 0x24)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GREEN  = RGBColor(0x1E, 0x7A, 0x33)
ORANGE = RGBColor(0xB4, 0x60, 0x10)
RED    = RGBColor(0xA3, 0x2A, 0x1F)

SERIF = "Georgia"
SANS  = "Calibri"

doc = Document()

# margenes
for s in doc.sections:
    s.top_margin = Inches(0.8); s.bottom_margin = Inches(0.8)
    s.left_margin = Inches(0.75); s.right_margin = Inches(0.75)

# estilo base
base = doc.styles["Normal"]
base.font.name = SANS
base.font.size = Pt(10)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)

def set_widths(table, widths):
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)

def cell_text(cell, text, bold=False, size=9, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size); r.font.color.rgb = color; r.font.name = SANS
    return p

def conf_color(c):
    if c.startswith("OK"): return GREEN
    if c.startswith("REV"): return ORANGE
    return RED

# ---------- portada / encabezado ----------
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = t.add_run("45 DIGITAL NOTICIAS")
r.font.name = SANS; r.font.size = Pt(9); r.font.color.rgb = AMBER; r.bold = True
r.font.all_caps = True

h = doc.add_paragraph()
rh = h.add_run("Directorio de difusión")
rh.font.name = SERIF; rh.font.size = Pt(26); rh.font.color.rgb = INK; rh.bold = True
h.paragraph_format.space_after = Pt(2)

sub = doc.add_paragraph()
rs = sub.add_run("Dashboard  ·  Inseguridad México · 32 estados")
rs.font.name = SERIF; rs.font.size = Pt(13); rs.font.color.rgb = AMBER; rs.italic = True
sub.paragraph_format.space_after = Pt(6)

# filete
fl = doc.add_paragraph()
pPr = fl._p.get_or_add_pPr()
pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "12")
bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "B47A10")
pbdr.append(bottom); pPr.append(pbdr)

meta = doc.add_paragraph()
rm = meta.add_run("Mapeo de correos públicos · 16 de junio de 2026 · uso interno de difusión")
rm.font.size = Pt(9); rm.font.color.rgb = GREY; rm.italic = True
meta.paragraph_format.space_after = Pt(10)

# leyenda de confianza
leg = doc.add_paragraph()
leg.add_run("Confianza del correo:  ").bold = True
def legrun(p, sym, txt, col):
    a = p.add_run(sym + " "); a.bold = True; a.font.color.rgb = col; a.font.size = Pt(9)
    b = p.add_run(txt + "    "); b.font.size = Pt(9); b.font.color.rgb = GREY
legrun(leg, "✓", "Verificado en sitio/columna oficial", GREEN)
legrun(leg, "!", "Confirmar antes de enviar", ORANGE)
legrun(leg, "x", "Sin correo (formulario/redes)", RED)
leg.paragraph_format.space_after = Pt(4)

rule = doc.add_paragraph()
rr = rule.add_run("Regla de uso: envío personalizado, en tandas, nunca lista masiva oculta. Reconfirmar los marcados «!» antes de mandar. Son contactos institucionales públicos.")
rr.font.size = Pt(8.5); rr.italic = True; rr.font.color.rgb = GREY
rule.paragraph_format.space_after = Pt(12)

# ---------- datos ----------
# (nombre, correo, confianza[OK/REV/NO], nota)
SECCIONES = [
 ("1 · Sociedad civil y observatorios de seguridad", [
   ("Observatorio Nacional Ciudadano (ONC)", "contacto@onc.org.mx", "OK",  "MÁXIMA PRIORIDAD. Incidencia delictiva con datos SESNSP; coordina la Red Nacional de Observatorios (33 locales)."),
   ("México Unido Contra la Delincuencia (MUCD)", "comunicacion@mucd.org.mx", "OK", "Buzón de Comunicación. OSC histórica en política de seguridad y datos delictivos."),
   ("Impunidad Cero", "contacto@impunidadcero.org", "OK", "Mide impunidad con datos duros (homicidios, denuncias)."),
   ("World Justice Project México (prensa)", "press@worldjusticeproject.mx", "OK", "Índices de Estado de Derecho e impunidad estatal. Buzón de prensa."),
   ("Intersecta", "contacto@intersecta.org", "OK", "Violencia con datos SESNSP/INEGI y perspectiva de género."),
   ("Elementa DDHH (México)", "info.mexico@elementaddhh.org", "OK", "Política de drogas y violencia, enfoque sociojurídico."),
   ("Reinserta", "contacto@reinserta.org", "OK", "Prevención del delito, sistema penitenciario y adolescentes."),
   ("TResearch (Carlos Penna, Dir.)", "carlos@tresearch.mx", "OK", "Estudios de opinión sobre homicidios/seguridad."),
   ("México Evalúa (Seguridad/Justicia)", "info@mexicoevalua.org", "REV", "ALTA PRIORIDAD. Programa de Seguridad y Justicia que analiza datos del SESNSP. Reconfirmar."),
   ("Causa en Común (M. E. Morera)", "donaciones@causaencomun.org.mx", "REV", "Datos de violencia, atrocidades y policías. Solo visible el de donativos; buscar editorial."),
   ("Data Cívica", "info@datacivica.org", "REV", "Datos de violencia y desaparición, metodología abierta. Muy afín. Reconfirmar."),
   ("CIDAC", "info@cidac.org", "REV", "Proyecto Justicia (medición de procuración/impartición). Reconfirmar."),
   ("Ethos Innovación en Políticas Públicas", "info@ethos.org.mx", "REV", "Seguridad y economía del crimen. Reconfirmar."),
   ("Lantia Intelligence (E. Guerrero)", "contacto@lantiaintelligence.com", "REV", "Violencia y crimen organizado con datos de fuentes abiertas."),
   ("Semáforo Delictivo", "(solo app / Facebook)", "NO", "Metodología afín, sin correo oficial. Contacto vía Facebook."),
 ]),
 ("2 · Academia y think tanks", [
   ("INSYDE — Inst. para la Seguridad y la Democracia", "info@insyde.org.mx", "OK", "Think tank nuclear en reforma policial y seguridad democrática."),
   ("ITAM — CESIG (Seguridad, Inteligencia, Gobernanza)", "cesig@itam.mx", "OK", "Centro académico específico de seguridad e inteligencia."),
   ("Ibero — Programa de Seguridad Ciudadana", "psc@ibero.mx", "OK", "Programa dedicado a seguridad ciudadana y «vía civil»."),
   ("CIDE — Div. de Estudios Jurídicos", "ximena.medellin@cide.edu / virginia.tovar@cide.edu", "OK", "Investigación sobre sistema penal y política de drogas."),
   ("CIDE — Programa de Política de Drogas", "programadepoliticadedrogas@gmail.com", "REV", "Programa eje en política de drogas y violencia (cuenta gmail; reconfirmar)."),
   ("El Colef (frontera, seguridad/migración)", "difusionuec@colef.mx / miginter@colef.mx", "OK", "Violencia, seguridad pública y DDHH en la frontera."),
   ("COLMEX — Centro de Estudios Sociológicos", "direccion.ces@colmex.mx", "OK", "Violencia y conflicto social desde la sociología."),
   ("Instituto Belisario Domínguez (Senado)", "apoyotecnico.ibd@senado.gob.mx", "OK", "Investigación estratégica del Senado; vincula al circuito legislativo."),
   ("Integralia Consultores", "contacto@integralia.com.mx", "OK", "Consultora de riesgo político/social."),
   ("UNAM — Inst. de Investigaciones Sociales", "iis@sociales.unam.mx", "OK", "Mayor centro de sociología de la UNAM; violencia y conflicto."),
   ("UNAM — CISAN (América del Norte)", "cisan@unam.mx", "OK", "Seguridad regional y relación con EE. UU."),
 ]),
 ("3 · Medios nacionales (opinión / redacción / seguridad)", [
   ("Carlos Loret de Mola (El Universal / Latinus)", "historiasreportero@gmail.com", "OK", "CONTACTO DE MÁXIMO VALOR. Él publica este correo al pie de su columna (la más leída del país) y dirige Latinus. Trabaja series de homicidio."),
   ("El Economista (redacción)", "redaccion@eleconomista.mx", "OK", "Único buzón editorial 100% verificado del lote."),
   ("Aristegui Noticias (redacción)", "redaccion@aristeguinoticias.com", "OK", "Buzón editorial general; sirve para opinión."),
   ("El Universal (investigación/seguridad)", "denuncia@eluniversal.com.mx", "OK", "Lo más cercano a un desk de investigación/seguridad."),
   ("El Universal (general/colaboraciones)", "contacto@eluniversal.com.mx", "OK", "Canal por defecto para envíos."),
   ("Quinto Elemento Lab (colaboraciones)", "quintoelab@gmail.com", "OK", "Periodismo de investigación; invita a colaborar. Afín a datos de violencia."),
   ("Infobae México (redacción)", "redaccion@infobae.com", "OK", "Buzón corporativo."),
   ("Expansión Política (editora)", "mibarra@grupoexpansion.com", "OK", "«Política» es el desk más cercano a seguridad/justicia. También ariadna.ortega@ y antonio.baranda@grupoexpansion.com."),
   ("Eje Central (redacción)", "comenta@ejecentral.com.mx", "OK", "Único buzón editorial público."),
   ("Latinus (colaboraciones)", "contacto@latinus.us", "OK", "Se nutre de colaboradores; vía para proponer columna."),
   ("Excélsior (comentarios)", "mvelazquezm@gimm.com.mx", "OK", "Único correo del directorio."),
   ("Milenio (general digital)", "digital@milenio.com", "OK", "Resto vía formulario con selector «Editorial»."),
   ("Animal Político (El Sabueso)", "elsabueso@animalpolitico.com", "REV", "Lo más cercano a justicia/seguridad que publican. Reconfirmar dominio."),
   ("Reforma (cartas / opinión)", "cartas@reforma.com", "REV", "Migra a formulario. Reconfirmar."),
   ("El Financiero (contacto web)", "contactoweb@elfinanciero.com.mx", "REV", "Aparece en directorios; reconfirmar."),
 ]),
 ("4 · Morelos y regional (caso profundo del dashboard)", [
   ("Observatorio de Seguridad Ciudadana (UAEM)", "obs-seguridad@obs-seguridad.org", "OK", "INTERLOCUTOR MÁS PERTINENTE del caso Morelos: incidencia delictiva."),
   ("IMIPE Morelos — Unidad de Transparencia", "imipe.ut@gmail.com", "OK", "El instituto sigue operando."),
   ("IMIPE — Comisionado Pdte. Hertino Avilés", "imipe.hertinoaviles@gmail.com", "OK", "Cabeza del órgano garante."),
   ("IMIPE — Socialización / Capacitación", "imipe.socializacion@gmail.com", "OK", "Área idónea para difundir un dashboard ciudadano. También imipe.redcapacitacion@gmail.com."),
   ("El Sol de Cuernavaca (OEM)", "tuopinion@elsoldecuernavaca.com.mx", "OK", "Cubre policiaca/seguridad."),
   ("El Regional del Sur", "contacto@elregional.com.mx", "OK", "Correo general de su página oficial."),
   ("Diario de Morelos", "redaccion@diariodemorelos.com", "REV", "Medio estatal de mayor peso; sección de seguridad. Reconfirmar. También socialmedia@diariodemorelos.com."),
   ("La Unión de Morelos", "uniondemorelos@gmail.com", "REV", "Reconfirmar. Dir.: alberto.sanchez@launion.com.mx; redacción: betty@launion.com.mx."),
   ("La Jornada Morelos", "redaccionjornadamorelos@gmail.com", "REV", "Reconfirmar. También correociudadano@lajornadamorelos.mx."),
   ("El Sol de Cuautla (OEM)", "elsoldecuautla@yahoo.com.mx", "REV", "Reconfirmar (registro SIC-Cultura)."),
   ("Ana María Salazar (periodista seguridad)", "anamariasalazar@gmail.com", "REV", "Columnista nacional de seguridad/justicia. Atribuido por fuentes secundarias; reconfirmar."),
   ("Jorge Medellín (EstadoMayor.mx)", "(solo X: @EstadoMayor_mx)", "NO", "Referente en seguridad/militar; sin correo público."),
 ]),
]

def add_section(titulo, filas):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(5)
    r = p.add_run(titulo)
    r.font.name = SANS; r.bold = True; r.font.size = Pt(12); r.font.color.rgb = INK

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    encabezados = ["Organización / Persona", "Correo", "Conf.", "Nota"]
    for i, e in enumerate(encabezados):
        shade(hdr[i], "0D1117")
        cell_text(hdr[i], e, bold=True, size=9, color=WHITE)
    for nombre, correo, conf, nota in filas:
        cells = table.add_row().cells
        cell_text(cells[0], nombre, bold=True, size=9, color=INK)
        cell_text(cells[1], correo, size=9, color=AMBER)
        sym = {"OK": "✓", "REV": "!", "NO": "x"}[conf]
        cell_text(cells[2], sym, bold=True, size=11, color=conf_color(conf), align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(cells[3], nota, size=8.5, color=GREY)
    set_widths(table, [2.15, 1.95, 0.45, 2.45])

for titulo, filas in SECCIONES:
    add_section(titulo, filas)

# ---------- plan de prioridad ----------
doc.add_paragraph().paragraph_format.space_after = Pt(2)
pp = doc.add_paragraph()
rp = pp.add_run("Orden de envío sugerido")
rp.font.name = SANS; rp.bold = True; rp.font.size = Pt(12); rp.font.color.rgb = INK
pp.paragraph_format.space_before = Pt(10)

pasos = [
 "Primer círculo (máximo encaje): Loret de Mola, ONC, México Evalúa, MUCD, Impunidad Cero, Intersecta, Data Cívica, WJP México.",
 "Caso Morelos: Observatorio UAEM + IMIPE (Socialización) + prensa estatal (El Sol de Cuernavaca, El Regional, Diario de Morelos).",
 "Academia / think tanks: INSYDE, CIDE-DEJ, ITAM-CESIG, Ibero-PSC, Belisario Domínguez.",
 "Medios para columna derivada: El Economista, Aristegui, El Universal, Expansión Política, Eje Central.",
]
for i, paso in enumerate(pasos, 1):
    lp = doc.add_paragraph()
    lp.paragraph_format.left_indent = Inches(0.2); lp.paragraph_format.space_after = Pt(3)
    rn = lp.add_run(f"{i}.  "); rn.bold = True; rn.font.color.rgb = AMBER; rn.font.size = Pt(9.5)
    rt = lp.add_run(paso); rt.font.size = Pt(9.5); rt.font.color.rgb = INK

# pendiente URL
nd = doc.add_paragraph()
nd.paragraph_format.space_before = Pt(8)
rnd = nd.add_run("Pendiente: insertar en el cuerpo del correo la URL pública exacta del dashboard (vive bajo 45digitalnoticias.github.io).")
rnd.italic = True; rnd.font.size = Pt(8.5); rnd.font.color.rgb = GREY

# firma
fm = doc.add_paragraph(); fm.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fm.paragraph_format.space_before = Pt(14)
rf = fm.add_run("45 Digital Noticias"); rf.italic = True; rf.font.size = Pt(9); rf.font.color.rgb = GREY

import os
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Directorio de difusion - Inseguridad Mexico.docx")
doc.save(out)
print("OK ->", out)
