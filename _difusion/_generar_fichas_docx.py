# -*- coding: utf-8 -*-
"""
Genera el Word de FICHAS INFORMATIVAS de cada contacto de difusion del dashboard
"Inseguridad Mexico". Una ficha por organizacion/persona, agrupadas por categoria.
Salida: Fichas de contactos - difusion Inseguridad Mexico.docx (en esta misma carpeta).
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

AMBER  = RGBColor(0xB4, 0x7A, 0x10)
GREY   = RGBColor(0x55, 0x5B, 0x63)
INK    = RGBColor(0x1A, 0x1E, 0x24)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GREEN  = RGBColor(0x1E, 0x7A, 0x33)
ORANGE = RGBColor(0xB4, 0x60, 0x10)
RED    = RGBColor(0xA3, 0x2A, 0x1F)
SANS  = "Calibri"

doc = Document()
for s in doc.sections:
    s.top_margin = Inches(0.7); s.bottom_margin = Inches(0.7)
    s.left_margin = Inches(0.85); s.right_margin = Inches(0.85)
doc.styles["Normal"].font.name = SANS
doc.styles["Normal"].font.size = Pt(10)

def add_runs(p, segs):
    for text, bold, color, size in segs:
        r = p.add_run(text); r.bold = bold
        r.font.color.rgb = color; r.font.size = Pt(size); r.font.name = SANS

def para(segs, justify=True, sa=4, sb=0, indent=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(sa); p.paragraph_format.space_before = Pt(sb)
    if indent: p.paragraph_format.left_indent = Inches(indent)
    add_runs(p, segs); return p

def hrule(color="D0D4DA", sz="6", sa=6):
    fl = doc.add_paragraph(); pPr = fl._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr"); b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), sz); b.set(qn("w:space"), "1"); b.set(qn("w:color"), color)
    pbdr.append(b); pPr.append(pbdr); fl.paragraph_format.space_after = Pt(sa)

def page_break():
    p = doc.add_paragraph(); run = p.add_run(); br = OxmlElement("w:br")
    br.set(qn("w:type"), "page"); run._r.append(br)

def conf_seg(conf):
    if conf == "OK":  return ("  ● verificado en fuente oficial", GREEN)
    if conf == "REV": return ("  ▲ dato por confirmar", ORANGE)
    return ("  ■ ojo: leer la nota", RED)

def ficha(nombre, conf, campos, fuente):
    seg, col = conf_seg(conf)
    para([(nombre, True, AMBER, 13), (seg, True, col, 8.5)], justify=False, sa=2, sb=6)
    for label, text in campos:
        lab_color = AMBER if label.startswith("Si te contactan") else GREY
        para([(label.upper() + ": ", True, lab_color, 8.5), (text, False, INK, 10)], sa=3)
    para([("Fuente: ", True, GREY, 8), (fuente, False, GREY, 8.5)], justify=False, sa=2)
    hrule(sa=8)

# ============== CONTENIDO ==============
# portada
para([("45 DIGITAL NOTICIAS", True, AMBER, 9.5)], justify=False, sa=2)
para([("Fichas de contactos · quién es quién", True, INK, 21)], justify=False, sa=2)
para([("Difusión del dashboard Inseguridad México · 16 de junio de 2026", True, AMBER, 11)], justify=False, sa=6)
hrule(color="B47A10", sz="10")
para([("Para que, si te responden, sepas con quién hablas: qué es la organización, quién la dirige, su línea en "
       "seguridad y cómo conversar con cada quien. Los nombres se verificaron en sitios oficiales; lo no confirmado "
       "va marcado. Leyenda: ● verificado · ▲ por confirmar · ■ ojo a la nota.", False, GREY, 9.5)], sa=10)
para([("ALERTA — IMIPE Morelos: fue EXTINGUIDO el 27-ene-2026 (decreto 1105) y sustituido por «Transparencia para el "
       "Pueblo de Morelos», órgano del Ejecutivo estatal. Ya no es el organismo autónomo. Ver su ficha.",
       True, RED, 9.5)], sa=6)

# ---------- 1. SOCIEDAD CIVIL ----------
page_break()
para([("1 · SOCIEDAD CIVIL Y OBSERVATORIOS", True, INK, 13)], justify=False, sa=8)

ficha("Observatorio Nacional Ciudadano (ONC)", "OK", [
 ("Tipo", "Observatorio ciudadano / OSC."),
 ("Qué es", "Organización que monitorea el desempeño de instituciones de seguridad y justicia y articula la Red Nacional de Observatorios. Vigilancia ciudadana basada en datos."),
 ("Figura clave", "Francisco Rivas Rodríguez, director general desde 2013 (también miembro del Consejo de Seguridad Nacional)."),
 ("Línea en seguridad/datos", "Reportes mensuales de incidencia con cifras del SESNSP por estado y municipio: homicidio, extorsión, feminicidio, desempeño de fiscalías."),
 ("Si te contactan", "Les interesa el dato duro georreferenciado y la metodología. Tono técnico-ciudadano. Ofréceles datos y colaboración metodológica, no solo difusión. Tu interlocutor #1."),
], "onc.org.mx")

ficha("México Evalúa", "OK", [
 ("Tipo", "Think tank de políticas públicas."),
 ("Qué es", "Centro que evalúa la acción de gobierno con evidencia: seguridad, justicia, finanzas públicas y anticorrupción."),
 ("Figura clave", "Armando Vargas, coordinador del Programa de Seguridad (dirección general institucional no confirmada en esta pasada; fundadora histórica: Edna Jaime)."),
 ("Línea en seguridad/datos", "Evaluación mensual de los 10 delitos de mayor impacto en las 32 entidades y +4,000 municipios. SU estimación de −8.6% es uno de los tres relojes de tu tablero."),
 ("Si te contactan", "Buscan rigor analítico y lectura de política pública, no nota roja. Tono institucional-académico. Reconoce que partes de su −8.6% y ofrece debate de hallazgos."),
], "mexicoevalua.org")

ficha("México Unido Contra la Delincuencia (MUCD)", "OK", [
 ("Tipo", "OSC."),
 ("Qué es", "AC de seguridad ciudadana, justicia y política de drogas; litigio estratégico y trabajo contra la militarización."),
 ("Figura clave", "Cristina Reyes Ortiz, directora general desde el 1-abr-2026 (releva a Lisa Sánchez); abogada, venía del litigio estratégico."),
 ("Línea en seguridad/datos", "Reforma policial, regulación de drogas, militarización, justicia penal y derechos; más incidencia y litigio que series de incidencia."),
 ("Si te contactan", "Les interesa el ángulo de derechos, drogas y militarización. Tono garantista. Más difusión e incidencia conjunta que intercambio de datos crudos."),
], "mucd.org.mx")

ficha("Impunidad Cero", "OK", [
 ("Tipo", "OSC / observatorio."),
 ("Qué es", "Mide y visibiliza la impunidad para empujar mejoras en procuración de justicia."),
 ("Figura clave", "Catalina Kühne Peimbert, directora ejecutiva (+20 años en servicio público federal)."),
 ("Línea en seguridad/datos", "Índice Estatal de Desempeño de Fiscalías, Índice de Impunidad y encuesta de percepciones; brecha entre delitos cometidos y castigados."),
 ("Si te contactan", "Buscan el ángulo de impunidad y desempeño de fiscalías. Tono técnico-jurídico. Ofréceles cruces sobre denuncia, sentencia y eficacia."),
], "impunidadcero.org")

ficha("World Justice Project México", "OK", [
 ("Tipo", "Think tank / capítulo México de organización internacional."),
 ("Qué es", "Brazo mexicano del WJP; mide la adherencia al Estado de Derecho desde la perspectiva ciudadana, con datos estatales."),
 ("Figura clave", "Titular del capítulo México no confirmado; cargos directivos visibles: Alejandro Ponce (investigación global) y Alejandro González Arreola (proyectos de Estado de Derecho)."),
 ("Línea en seguridad/datos", "Índice de Estado de Derecho en México (IEDMX), con cifras estatales en 8 factores (orden, seguridad, justicia penal); Índice Global."),
 ("Si te contactan", "Les interesa el dato comparado y la metodología. Tono académico-internacional. Cita bien sus índices y ofrece colaboración de datos."),
], "worldjusticeproject.mx")

ficha("Intersecta", "OK", [
 ("Tipo", "OSC feminista / observatorio de datos."),
 ("Qué es", "Organización feminista que combate la discriminación con análisis de datos, con foco en violencia de género."),
 ("Figura clave", "Estefanía Vela Barba, directora ejecutiva; Adriana G. Ortega coordina el área de datos."),
 ("Línea en seguridad/datos", "Violencia de género, feminicidio, violencia institucional; lectura crítica de cifras oficiales de violencia contra las mujeres."),
 ("Si te contactan", "Les interesa la perspectiva de género e interseccional sobre los datos. Tono feminista-técnico. Ofrece desagregación por sexo y colaboración analítica."),
], "intersecta.org")

ficha("Data Cívica", "REV", [
 ("Tipo", "OSC / observatorio de datos."),
 ("Qué es", "(2015) Usa datos y tecnología para defender DDHH: violencia, feminicidio, desaparición y memoria."),
 ("Figura clave", "Mónica Meltis, directora ejecutiva desde 2017 (nombre verificado en fuentes serias; el sitio bloqueó el acceso directo)."),
 ("Línea en seguridad/datos", "Bases y herramientas sobre desaparecidos, feminicidio y violaciones graves; análisis crítico de cifras oficiales."),
 ("Si te contactan", "Les interesa el dato abierto y el rigor técnico. Tono de ciencia de datos con enfoque DDHH. Ofréceles datos estructurados y colaboración técnica."),
], "datacivica.org")

ficha("Causa en Común", "OK", [
 ("Tipo", "OSC."),
 ("Qué es", "(2010) Ciudadanía, transparencia y rendición de cuentas, con foco fuerte en la dignificación policial."),
 ("Figura clave", "María Elena Morera Mitre, presidenta y fundadora (vocal y activa en 2026)."),
 ("Línea en seguridad/datos", "Monitoreo de policías, condiciones laborales policiales, registro de masacres y atrocidades; crítica a la militarización."),
 ("Si te contactan", "Les interesa el ángulo policial y la rendición de cuentas. Tono frontal, crítico al poder. Ofrece datos sobre fuerzas de seguridad y difusión conjunta."),
], "causaencomun.org.mx")

ficha("Elementa DDHH", "OK", [
 ("Tipo", "OSC / consultoría en derechos humanos."),
 ("Qué es", "(2014) Oficinas en México y Colombia; usa el derecho para garantizar DDHH; política de drogas y justicia transicional."),
 ("Figura clave", "Renata Demichelis Ávila (directora, México); Paula Aguirre Ospina (Colombia)."),
 ("Línea en seguridad/datos", "Política de drogas con enfoque DDHH, verdad/justicia/reparación; investigación técnico-jurídica e incidencia."),
 ("Si te contactan", "Les interesa el cruce drogas-DDHH y la dimensión regional. Tono jurídico-garantista. Más marco normativo que datos crudos."),
], "elementaddhh.org")

ficha("Reinserta", "OK", [
 ("Tipo", "OSC."),
 ("Qué es", "(2013) Atiende a niñez y adolescencia en contacto con la violencia y el sistema penal, con modelos psicosociales y psico-jurídicos."),
 ("Figura clave", "Saskia Niño de Rivera, cofundadora y presidenta (cofundó con Mercedes Castañeda)."),
 ("Línea en seguridad/datos", "Niñez/adolescencia en prisión y violencia, reinserción, maternidad en reclusión."),
 ("Si te contactan", "Les interesa el ángulo de infancia, cárceles y reinserción, no la macro-incidencia. Tono humano y de denuncia. Ofrece datos de menores y sistema penal."),
], "reinserta.org")

ficha("TResearch (Carlos Penna)", "OK", [
 ("Tipo", "Consultora de investigación y opinión."),
 ("Qué es", "Sistematiza estadística delictiva pública en informes periódicos de seguridad."),
 ("Figura clave", "Carlos Penna Charolet, director."),
 ("Línea en seguridad/datos", "«La guerra en números» y «México en números»: homicidios, violencia política y familiar, desapariciones, robo de autos, ranking de municipios; series ENSU/INEGI."),
 ("Si te contactan", "Les interesa el dato cuantitativo comparable y la serie histórica. Tono de consultora de datos (ICC/ESOMAR). Intercambio técnico de cifras."),
], "tresearch.mx")

ficha("CIDAC", "REV", [
 ("Tipo", "Think tank."),
 ("Qué es", "Centro independiente que propone políticas en desarrollo, justicia y seguridad para el Estado de Derecho."),
 ("Figura clave", "Verónica Baz (directora general) y Luis Rubio (presidente), vía fuente secundaria; no confirmado en sitio para 2026."),
 ("Línea en seguridad/datos", "Histórico Índice Delictivo (delitos de alto impacto por entidad); implementación del sistema de justicia penal acusatorio."),
 ("Si te contactan", "Les interesa la política pública basada en evidencia y el sistema de justicia. Tono técnico-institucional. Ofrece datos y discusión de hallazgos."),
], "cidac.org")

ficha("Ethos Innovación en Políticas Públicas", "REV", [
 ("Tipo", "Think tank."),
 ("Qué es", "Recomendaciones de política pública en finanzas, anticorrupción, inclusión y desarrollo sostenible."),
 ("Figura clave", "Liliana Alvarado, directora ejecutiva (ex SHCP; ITAM/LSE)."),
 ("Línea en seguridad/datos", "Seguridad no es su eje; fuerte en anticorrupción y finanzas. Conexión indirecta: gasto público, corrupción, economía del crimen."),
 ("Si te contactan", "Les interesa el ángulo del dinero detrás de la inseguridad. Tono de política pública/economía. Difusión y datos de presupuesto más que incidencia."),
], "ethos.org.mx")

ficha("Lantia Intelligence (Eduardo Guerrero)", "REV", [
 ("Tipo", "Consultora privada de inteligencia."),
 ("Qué es", "Inteligencia y ciencia de datos sobre seguridad y crimen organizado."),
 ("Figura clave", "Eduardo Guerrero Gutiérrez, socio fundador y director (rol confirmado en múltiples fuentes)."),
 ("Línea en seguridad/datos", "Cartografía de grupos criminales, homicidios, extorsión, desapariciones; fragmentación de cárteles desde 2008."),
 ("Si te contactan", "Análisis de crimen organizado de alto nivel. Tono de consultoría premium. Más intercambio analítico que difusión abierta; es firma privada."),
], "lantiaintelligence.com")

# ---------- 2. ACADEMIA ----------
page_break()
para([("2 · ACADEMIA Y THINK TANKS", True, INK, 13)], justify=False, sa=8)

ficha("INSYDE — Inst. para la Seguridad y la Democracia", "OK", [
 ("Tipo", "Think tank / OSC."),
 ("Qué es", "(2003) ONG transdisciplinaria para la reforma policial y la convivencia democrática."),
 ("Figura clave", "Sophie Anaya Levesque (presidenta del Consejo); Miguel Garza Flores (director ejecutivo)."),
 ("Línea en seguridad/datos", "Reforma policial democrática, justicia penal y DDHH, movilidad humana, violencia y medios. Seguridad ciudadana, no punitivista."),
 ("Si te contactan", "Les interesa el control civil de la policía y el enfoque de derechos; crítico del modelo militarizado. Ofrece datos georreferenciados y colaboración; valoran evidencia."),
], "insyde.org.mx")

ficha("ITAM — CESIG", "OK", [
 ("Tipo", "Centro de investigación universitario (ITAM)."),
 ("Qué es", "Conocimiento sobre seguridad, inteligencia estratégica y gobernanza, con evaluación de políticas."),
 ("Figura clave", "Codirectores Vidal Romero y Jorge Tello Peón."),
 ("Línea en seguridad/datos", "Seguridad, inteligencia y gobernanza; opera desde 2013 el Seminario sobre Seguridad, Ciudadanía y Violencia."),
 ("Si te contactan", "Perfil técnico-cuantitativo, cercano a decisores e inteligencia/Estado. Tono académico riguroso. Ofrece metodología y datos abiertos; aprecian el rigor estadístico."),
], "cesig.itam.mx")

ficha("Ibero — Programa de Seguridad Ciudadana", "OK", [
 ("Tipo", "Programa universitario (Ibero CDMX)."),
 ("Qué es", "(2018) Incidencia y formación en seguridad ciudadana ligada a DDHH; maestría propia."),
 ("Figura clave", "Ernesto López Portillo Vargas (coordinador)."),
 ("Línea en seguridad/datos", "Alternativas a la militarización, reforma del SNSP, Guardia Nacional, control civil, prevención."),
 ("Si te contactan", "Les interesa la crítica a la militarización y la agenda civil. Fuerte presencia mediática. Colaboración y difusión mutua; López Portillo amplifica datos."),
], "ibero.mx")

ficha("CIDE — División de Estudios Jurídicos", "OK", [
 ("Tipo", "Centro de investigación universitario (CIDE)."),
 ("Qué es", "Estudios jurídicos empíricos e interdisciplinarios (derecho, política, análisis social)."),
 ("Figura clave", "Dra. Ximena Medellín Urquiaga (jefa de la División)."),
 ("Línea en seguridad/datos", "Criminología, sistema de justicia y derecho penal, seguridad pública, estudios jurídicos empíricos."),
 ("Si te contactan", "Perfil jurídico-empírico, muy afín a tu propio perfil de abogado. Tono académico, basado en evidencia. Ofrece microdatos de justicia/seguridad y colaboración."),
], "cide.edu/division_dej")

ficha("CIDE — Programa de Política de Drogas", "REV", [
 ("Tipo", "Programa de investigación universitario (CIDE)."),
 ("Qué es", "Estudia el fenómeno y las políticas de drogas en México y América Latina."),
 ("Figura clave", "Laura Atuesta (coordinadora) no confirmada en sitio; fundador previo: Alejandro Madrazo."),
 ("Línea en seguridad/datos", "Política de drogas, relación prohibición-violencia, datos de letalidad y operativos, regulación de cannabis."),
 ("Si te contactan", "Les interesa la evidencia sobre violencia ligada a la guerra contra las drogas. Tono académico crítico de la prohibición. Comparte datos desagregados."),
], "politicadedrogas.org")

ficha("El Colegio de la Frontera Norte (El Colef)", "OK", [
 ("Tipo", "Centro público de investigación (SECIHTI)."),
 ("Qué es", "Conocimiento científico sobre la frontera México-EE.UU.; posgrados y asesoría regional."),
 ("Figura clave", "Dr. Víctor Alejandro Espinoza Valle (presidente); recotejar por posible relevo de gestión."),
 ("Línea en seguridad/datos", "Migración, seguridad humana, violencia fronteriza, narcotráfico transfronterizo; opera la EMIF."),
 ("Si te contactan", "Enfoque regional y migratorio. Tono académico de datos de campo. Cruza con la dimensión fronteriza y ofrece datos de los estados del norte."),
], "colef.mx")

ficha("COLMEX — Centro de Estudios Sociológicos", "OK", [
 ("Tipo", "Centro de investigación universitario (COLMEX)."),
 ("Qué es", "Centro de sociología del Colegio de México; investigación social de alto nivel; revista Estudios Sociológicos."),
 ("Figura clave", "Gustavo Adolfo Urbina Cortés (director interino); presidenta de COLMEX: Ana Covarrubias (desde 2025)."),
 ("Línea en seguridad/datos", "Sociología de la violencia, desigualdad, demografía, estudios urbanos; análisis estructural más que coyuntura."),
 ("Si te contactan", "Perfil sociológico de prestigio. Tono conceptual y riguroso. Ofrece datos para análisis estructural; orientado a publicación arbitrada más que difusión rápida."),
], "ces.colmex.mx")

ficha("Instituto Belisario Domínguez (Senado)", "REV", [
 ("Tipo", "Órgano técnico del Senado de la República."),
 ("Qué es", "Produce investigación estratégica y análisis legislativo para apoyar decisiones del Congreso."),
 ("Figura clave", "Sen. Ignacio Mier Velazco (presidente de la Junta Directiva); director general por nombre no confirmado en sitio."),
 ("Línea en seguridad/datos", "La Dirección General de Investigación Estratégica cubre DDHH, seguridad y justicia, con estudios comparados y de implementación."),
 ("Si te contactan", "Instancia oficial del Legislativo; su material es insumo para legisladores. Tono institucional y prudente. Difusión y datos citables; cuida el blindaje y la atribución."),
], "ibd.senado.gob.mx")

ficha("Integralia Consultores", "OK", [
 ("Tipo", "Think tank / consultora de riesgo político."),
 ("Qué es", "(2009) Análisis prospectivo de riesgo político y de seguridad, para empresas e inversionistas."),
 ("Figura clave", "Luis Carlos Ugalde (fundador y director; expresidente del IFE 2003-2007)."),
 ("Línea en seguridad/datos", "Riesgo político-institucional, seguridad y clima de negocios; informe anual «Diez riesgos políticos para México»."),
 ("Si te contactan", "Audiencia corporativa y financiera, no académica. Tono ejecutivo y prospectivo. Posiciona el tablero como insumo de riesgo país; útil para difusión de alto nivel."),
], "integralia.com.mx")

ficha("UNAM — Instituto de Investigaciones Sociales", "OK", [
 ("Tipo", "Instituto de investigación universitario (UNAM)."),
 ("Qué es", "Principal instituto de ciencias sociales de la UNAM; grandes problemas sociales de México."),
 ("Figura clave", "Dra. Marcela Amaro Rosales (directora)."),
 ("Línea en seguridad/datos", "Violencia social, tráfico de drogas e inseguridad; desplazamiento forzado interno."),
 ("Si te contactan", "Perfil académico-estructural amplio. Tono universitario riguroso. Ofrece datos para análisis de violencia y desplazamiento; abiertos a colaboración de mediano plazo."),
], "iis.unam.mx")

ficha("UNAM — CISAN (América del Norte)", "OK", [
 ("Tipo", "Centro de investigación universitario (UNAM)."),
 ("Qué es", "Estudio multidisciplinario de América del Norte (México, EE.UU., Canadá)."),
 ("Figura clave", "Dr. Juan Carlos Barrón Pastor (director, periodo 2025-2029)."),
 ("Línea en seguridad/datos", "Relación de seguridad México-EE.UU., narcotráfico transnacional, migración y fronteras."),
 ("Si te contactan", "Enfoque geopolítico y de relaciones internacionales, muy afín a tu lectura sistémica. Tono pluralista. Encuadra tus datos en lo bilateral/transnacional."),
], "cisan.unam.mx")

# ---------- 3. MEDIOS ----------
page_break()
para([("3 · MEDIOS NACIONALES", True, INK, 13)], justify=False, sa=8)

ficha("Carlos Loret de Mola", "OK", [
 ("Tipo", "Periodista / columnista."),
 ("Qué es", "Uno de los periodistas más conocidos del país (+24 años). Condujo «Primero Noticias» y «Despierta» en Televisa; hoy columnista de El Universal, conduce en W Radio y fundó/dirige Latinus (2020)."),
 ("Perfil", "Reportero y excorresponsal de guerra; en fuentes públicas se le identifica como una de las voces más críticas del obradorismo, confrontado abiertamente por el gobierno de AMLO."),
 ("Cobertura de seguridad", "Sí; su sello son investigaciones de corrupción, narco y abuso de poder."),
 ("Si te contactan", "Le interesan datos duros, documentos y exclusivas con ángulo de poder. Tono directo. Ofrece el tablero como insumo verificable; una entrevista en Latinus/W Radio amplifica mucho pero te ata a su línea. Tu contacto de mayor alcance."),
], "eluniversal.com.mx (columna)")

ficha("El Economista", "OK", [
 ("Tipo", "Diario de economía y negocios (impreso y digital)."),
 ("Qué es", "(1988) Referencia financiera para tomadores de decisión; propiedad del grupo Nacer Global."),
 ("Perfil", "Se presenta como independiente y técnico-económico, más que de confrontación política."),
 ("Cobertura de seguridad", "Limitada como eje propio; aborda seguridad cuando cruza con economía (costo de la violencia, riesgo país, huachicol)."),
 ("Si te contactan", "Les interesa el ángulo económico de la inseguridad: cifras, impacto en negocios, series. Tono sobrio y cuantitativo. Ofrece el tablero como fuente de datos."),
], "eleconomista.mx")

ficha("Aristegui Noticias", "OK", [
 ("Tipo", "Portal digital."),
 ("Qué es", "(2012) Liderado por Carmen Aristegui; de los sitios de mayor alcance y seguidores en México."),
 ("Perfil", "Periodismo de investigación e independiente («La Casa Blanca de Peña Nieto»); crítico del poder en turno, sea cual sea el partido."),
 ("Cobertura de seguridad", "Sí; corrupción, DDHH, abuso de poder y violencia estatal."),
 ("Si te contactan", "Buscan exclusivas documentadas y de interés público con rigor. Tono serio. Ofrece datos verificables y, si hay hallazgo fuerte, una columna o entrevista (alto alcance, exigente con la verificación)."),
], "aristeguinoticias.com")

ficha("El Universal", "OK", [
 ("Tipo", "Diario nacional (impreso y digital)."),
 ("Qué es", "De los diarios más antiguos y de mayor circulación; ediciones regionales, El Gráfico y agencia."),
 ("Perfil", "Propiedad de la familia Ealy Ortiz desde 1969; se declara plural y aloja firmas de muy distinto signo (incluido Loret); en análisis de medios se le ha asociado a una línea cercana al poder."),
 ("Cobertura de seguridad", "Sí, sección nacional/seguridad robusta y diaria."),
 ("Si te contactan", "Buen hogar para la columna (gran difusión) y para que sus reporteros de seguridad usen el tablero como fuente. Tono institucional."),
], "eluniversal.com.mx")

ficha("Quinto Elemento Lab", "OK", [
 ("Tipo", "Laboratorio de periodismo de investigación (sin fines de lucro)."),
 ("Qué es", "(2017) Reportajes propios de largo aliento y capacitación; premios Breach-Valdez y Gabo."),
 ("Perfil", "Fundado por Alejandra Xanic, Marcela Turati e Ignacio Rodríguez Reyna; eje en desigualdad, impunidad y corrupción, con periodismo de datos."),
 ("Cobertura de seguridad", "Muy fuerte; Marcela Turati es referente en desapariciones y fosas."),
 ("Si te contactan", "Les interesan investigaciones profundas y bases de datos, no la nota del día. Tono metodológico. Ofrece el tablero como dataset y una posible colaboración, no una columna de opinión."),
], "quintoelab.org")

ficha("Infobae México", "OK", [
 ("Tipo", "Portal digital."),
 ("Qué es", "Edición mexicana del medio nativo argentino Infobae (en México desde 2018); uno de los más leídos, altísimo tráfico."),
 ("Perfil", "Modelo de gran volumen y velocidad, sin muro de pago; fuentes públicas lo describen de tendencia conservadora por la trayectoria de su dueño (Daniel Hadad)."),
 ("Cobertura de seguridad", "Amplia y diaria; sección de seguridad/narco de mucho volumen, orientada a tráfico."),
 ("Si te contactan", "Buscan inmediatez, datos llamativos y material visual. Tono ágil. Ofrece el tablero con un gancho noticioso claro y titulares fuertes pero verificables."),
], "infobae.com")

ficha("Expansión Política (Grupo Expansión)", "OK", [
 ("Tipo", "Revista / portal digital (vertical de política)."),
 ("Qué es", "Expansión es la revista insignia de negocios (1969); «Expansión Política» es su sección de política y poder."),
 ("Perfil", "Se define independiente y rigurosa; perfil técnico-analítico dirigido a audiencias de decisión."),
 ("Cobertura de seguridad", "Moderada; seguridad sobre todo por su impacto institucional, político y económico."),
 ("Si te contactan", "Les interesa el ángulo de política pública y poder con datos. Tono analítico y sobrio. Ofrece el tablero más un análisis de implicaciones de política de seguridad."),
], "expansion.mx")

ficha("Eje Central", "OK", [
 ("Tipo", "Portal digital de análisis político."),
 ("Qué es", "(2009) Fundado por Raymundo Riva Palacio; dirigido a líderes de opinión, fuerte en columnas."),
 ("Perfil", "Riva Palacio es de las plumas más influyentes; su columna «Estrictamente Personal» suele ser crítica del gobierno en turno. Muy leído en el círculo del poder."),
 ("Cobertura de seguridad", "Sí, desde la inteligencia, la política de seguridad y el poder."),
 ("Si te contactan", "Les interesa el análisis fino y el dato que reconfigura una lectura política. Tono de élite informada. Encaja muy bien una columna con tesis interpretativa."),
], "ejecentral.com.mx")

ficha("Latinus", "REV", [
 ("Tipo", "Portal digital."),
 ("Qué es", "(2020) Plataforma fundada por Carlos Loret de Mola, gran alcance en redes; conocida por investigaciones de corrupción."),
 ("Perfil", "En fuentes públicas se le identifica como editorialmente crítico de la 4T. Su financiamiento ha sido objeto de reportes y disputa pública; tratar ese tema como controversia atribuida, no como hecho propio."),
 ("Cobertura de seguridad", "Sí; corrupción, narco y abuso de poder son su marca."),
 ("Si te contactan", "Les interesan exclusivas y documentos con impacto. Tono confrontativo con el poder. Ofrece datos verificables; entrevista de alto alcance, atada a su línea."),
], "latinus.us")

ficha("Excélsior", "OK", [
 ("Tipo", "Diario nacional (impreso y digital)."),
 ("Qué es", "Diario histórico (cien años); desde 2006 propiedad de Grupo Imagen (Vázquez Raña / Vázquez Aldir), con TV y radio."),
 ("Perfil", "Se presenta plural «sin sensacionalismo»; por ser un conglomerado con intereses amplios, conviene leer su línea como cercana al establishment económico."),
 ("Cobertura de seguridad", "Sí, diaria, reforzada por la sinergia con Imagen Televisión."),
 ("Si te contactan", "Buen vehículo para la columna y para alimentar su mesa de seguridad. Tono formal."),
], "excelsior.com.mx")

ficha("Milenio", "OK", [
 ("Tipo", "Grupo multimedia (diario, web, TV y radio)."),
 ("Qué es", "Propiedad de Grupo Multimedios; edición nacional y ocho regionales, gran alcance combinado."),
 ("Perfil", "Se declara de plena independencia editorial; medio masivo y plural, fuerte en el norte, aloja firmas diversas."),
 ("Cobertura de seguridad", "Sí, robusta; nacional y regional, incluida su pantalla de TV."),
 ("Si te contactan", "Les interesa material ágil y con datos, y posibilidad de TV. Tono masivo. Ofrece el tablero (visual, citable) y una columna; el ángulo regional es un plus."),
], "milenio.com")

ficha("Animal Político (y El Sabueso)", "OK", [
 ("Tipo", "Portal digital / unidad de verificación."),
 ("Qué es", "(2010) Medio independiente de periodismo de datos e investigación. El Sabueso (2015) es su fact-checking, certificado por la IFCN."),
 ("Perfil", "Dirigido por Daniel Moreno; rigor, datos y rendición de cuentas, sin alineación partidista declarada. Dos veces Premio Nacional de Periodismo."),
 ("Cobertura de seguridad", "Muy fuerte; tiene Narcodata (base sobre crimen organizado) y cobertura por datos."),
 ("Si te contactan", "Les interesan datos verificables y metodología transparente; El Sabueso podría cotejar tus cifras. Ofrece el tablero como dataset documentado y blinda cada número."),
], "animalpolitico.com")

ficha("Reforma / Grupo Reforma", "OK", [
 ("Tipo", "Diario de prestigio (impreso y digital, con muro de pago)."),
 ("Qué es", "(1993) Cabecera de Grupo Reforma (El Norte, Mural); dirigido a élites, modelo de suscripción."),
 ("Perfil", "Propiedad de la familia Junco; conocido por su independencia frente al poder y por roces públicos con el gobierno de AMLO. Estándar exigente."),
 ("Cobertura de seguridad", "De las más fuertes del país; investigación de seguridad y crimen organizado de primer nivel."),
 ("Si te contactan", "Les interesan datos sólidos y rigor; verifican con dureza. Tono formal y crítico. Ofrece el tablero como fuente para sus reporteros; su espacio de opinión es selectivo."),
], "reforma.com")

ficha("El Financiero", "OK", [
 ("Tipo", "Diario de finanzas (impreso, digital y TV)."),
 ("Qué es", "(1981) Propiedad de Grupo Multimedia Lauman; desde 2014 alianza El Financiero-Bloomberg."),
 ("Perfil", "Análisis y contextualización económica y empresarial; perfil técnico-financiero, con respaldo de Bloomberg."),
 ("Cobertura de seguridad", "Moderada; seguridad por su impacto económico (riesgo, inversión, extorsión a empresas)."),
 ("Si te contactan", "Les interesa el ángulo económico-financiero con datos duros. Tono sobrio y analítico. Ofrece el tablero con lectura de impacto económico."),
], "elfinanciero.com.mx")

# ---------- 4. MORELOS Y REGIONAL ----------
page_break()
para([("4 · MORELOS Y REGIONAL", True, INK, 13)], justify=False, sa=8)

ficha("Observatorio de Seguridad Ciudadana y Cohesión Social (UAEM)", "REV", [
 ("Tipo", "Observatorio académico (universitario)."),
 ("Qué es", "Plataforma de la UAEM que analiza la inseguridad en la zona metropolitana de Cuernavaca cruzando análisis criminal con variables sociales y espaciales."),
 ("Figura clave", "Alfonso Valenzuela Aguilera (director); por confirmar producción vigente (el sitio luce poco actualizado, hacia 2019-2020)."),
 ("Perfil / línea", "Académico, basado en evidencia, no partidista."),
 ("Si te contactan", "Interlocutor natural para metodología y validación de datos del caso Morelos. Tono técnico-académico. Ofrece colaboración (cotejo de indicadores), no solo difusión."),
], "obs-seguridad.org")

ficha("IMIPE Morelos — EXTINGUIDO (leer)", "NO", [
 ("Tipo", "Órgano de transparencia (ya extinto)."),
 ("Qué es", "Era el órgano autónomo garante del acceso a la información en Morelos. Extinguido por el decreto 1105 (Periódico Oficial «Tierra y Libertad», 27-ene-2026); ya no atiende solicitudes."),
 ("Quién lo sustituye", "«Transparencia para el Pueblo de Morelos», órgano desconcentrado de la Secretaría Anticorrupción y Buen Gobierno (encabezada por Alejandra Pani Barragán); el área la lleva Alejandra Fernández Hernández."),
 ("Perfil / línea", "Tema sensible: la disolución del autónomo y su traspaso al Ejecutivo es objeto de debate público sobre independencia. Toca tu línea editorial del desmantelamiento."),
 ("Si te contactan", "Si responde el órgano sucesor (gubernamental), trátalo con cautela y atribución. Usa tu tablero como palanca para pedir datos y documentar opacidad, NO como aval. No lo trates como el IMIPE autónomo."),
], "decreto 1105, Periódico Oficial Morelos (27-ene-2026)")

ficha("El Sol de Cuernavaca (OEM)", "REV", [
 ("Tipo", "Diario estatal (Organización Editorial Mexicana)."),
 ("Qué es", "Periódico de Cuernavaca de la mayor cadena de diarios del país; local de Morelos más contenido nacional de la red OEM."),
 ("Figura clave", "OEM la preside Paquita Ramos de Vázquez; dirección local específica no confirmada."),
 ("Perfil / línea", "Diario comercial de amplio alcance, línea institucional y de servicio; amplia cobertura de seguridad y nota local."),
 ("Si te contactan", "Buen vehículo de difusión por volumen y red nacional. Tono noticioso. Ofrece el caso Morelos del tablero con datos listos para nota."),
], "oem.com.mx/elsoldecuernavaca")

ficha("El Regional del Sur", "REV", [
 ("Tipo", "Diario estatal (Grupo Editorial Tlahuica)."),
 ("Qué es", "Matutino de Cuernavaca (1988); cobertura local de Morelos, activo en seguridad y política estatal."),
 ("Figura clave", "Eolo Ernesto Pacheco Rodríguez (director); Miguel Ángel Provisor (jefe de información)."),
 ("Perfil / línea", "Se asume plural e independiente; ha dado cobertura crítica al tema IMIPE/transparencia."),
 ("Si te contactan", "Receptivo a temas estatales y de seguridad. Tono local. Ofrece el caso Morelos y, si hay interés, una colaboración de columna/análisis."),
], "elregional.com.mx")

ficha("Diario de Morelos", "REV", [
 ("Tipo", "Diario estatal (Grupo Braca de Comunicación)."),
 ("Qué es", "Periódico de Cuernavaca (1978), de mayor presencia local en Morelos; impreso y digital."),
 ("Figura clave", "Familia Bracamontes (fundado por Federico Bracamontes Gálvez); titular actual al 2026 no confirmado."),
 ("Perfil / línea", "Se define plural y de servicio a la comunidad; referencia regional establecida."),
 ("Si te contactan", "Alcance local fuerte para difusión. Tono noticioso de servicio. Ofrece el caso Morelos con datos verificables."),
], "diariodemorelos.com")

ficha("La Unión de Morelos", "REV", [
 ("Tipo", "Diario estatal."),
 ("Qué es", "Diario de Cuernavaca de circulación estatal; política, sociedad y seguridad locales."),
 ("Figura clave", "Mario Estrada Elizondo (director, según registro de Cultura; verificar vigencia)."),
 ("Perfil / línea", "Diario local de referencia con amplia cobertura de política estatal; cubrió el nuevo modelo de transparencia."),
 ("Si te contactan", "Útil para difusión y agenda estatal. Tono institucional-noticioso. Ofrece el caso Morelos como dato duro para nota."),
], "launion.com.mx")

ficha("La Jornada Morelos", "REV", [
 ("Tipo", "Diario estatal (marca La Jornada)."),
 ("Qué es", "Edición morelense ligada a La Jornada; diario de Cuernavaca con cobertura local."),
 ("Figura clave", "Enrique Balp Díaz (director); Germán R. Muñoz G. (editor en jefe); marca La Jornada: Carmen Lira Saade."),
 ("Perfil / línea", "Línea progresista/de izquierda; atención a derechos, sociedad civil y crítica al poder; afín a coberturas de transparencia con enfoque social."),
 ("Si te contactan", "Receptivo a la lectura crítica y de derechos. Tono analítico. Buen espacio para colaboración o columna con énfasis sistémico."),
], "lajornadamorelos.mx")

ficha("El Sol de Cuautla (OEM)", "REV", [
 ("Tipo", "Diario regional (OEM)."),
 ("Qué es", "Diario de Cuautla de la red OEM; cubre la zona oriente de Morelos."),
 ("Figura clave", "Dirección local no confirmada al 2026 (registros previos citan a Daniel Martínez Castellanos); corporativo OEM: Paquita Ramos de Vázquez."),
 ("Perfil / línea", "Diario comercial de la red OEM, línea institucional; fuerte en nota local del oriente morelense."),
 ("Si te contactan", "Ideal para difusión enfocada al oriente. Tono noticioso. Ofrece el desglose por municipio del tablero (Cuautla suele ser punto caliente)."),
], "oem.com.mx/elsoldecuautla")

ficha("Ana María Salazar", "OK", [
 ("Tipo", "Periodista / analista de seguridad nacional."),
 ("Qué es", "Analista con +30 años en seguridad y política; abogada por Harvard Law, ex agregada judicial de la Embajada de EE.UU. en Bogotá y exfuncionaria del Pentágono. Dirige Grupo Salazar Slack."),
 ("Perfil / línea", "Enfoque técnico en seguridad nacional/pública y relación México-EE.UU. Columna «Análisis sin Fronteras» (desde 2001) y conducción en radio/TV. Perfil analítico, no militante."),
 ("Si te contactan", "Interlocutora de alto nivel: le interesan datos rigurosos y comparados. Tono técnico y sobrio. Ofrece el tablero como evidencia metodológicamente sólida (fuente, periodo, comparativo), más que como denuncia; posible cita experta o entrevista."),
], "elfinanciero.com.mx (columna)")

# ---------- cierre ----------
page_break()
para([("NOTAS DE VERIFICACIÓN", True, AMBER, 12)], justify=False, sa=8)
for t in [
 "Los nombres marcados ● se verificaron en el sitio oficial de la organización. Los ▲ provienen de fuente secundaria seria o de directorios que no rinden el dato al 2026; confírmalos al primer contacto.",
 "IMIPE Morelos: el hallazgo más sensible. El organismo autónomo fue extinguido (decreto 1105, 27-ene-2026). Cualquier respuesta «del IMIPE» vendría hoy del órgano sucesor del Ejecutivo, no del autónomo. Ajusta los correos a IMIPE en consecuencia.",
 "Medios de investigación/datos (Quinto Elemento, Animal Político/El Sabueso, Reforma): verifican con dureza. Con ellos, el tablero entra como dataset documentado y cada cifra debe resistir el cotejo.",
 "Ningún nombre fue inventado: donde no se pudo confirmar, se dijo expresamente.",
]:
    para([("• ", True, AMBER, 10), (t, False, INK, 9.5)], sa=4)
para([("45 Digital Noticias · Sergio Rubén Valdespín Oseguera", True, GREY, 9)], justify=False, sb=10)

out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Fichas de contactos - difusion Inseguridad Mexico.docx")
doc.save(out)
print("OK ->", out)
