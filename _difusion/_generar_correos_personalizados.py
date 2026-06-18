# -*- coding: utf-8 -*-
"""
Genera el Word de CORREOS PERSONALIZADOS de difusion del dashboard "Inseguridad Mexico".
Un correo por pagina, listo para copiar/pegar. Cuerpo de datos compartido + personalizacion
(saludo, asunto, parrafo "por que tu", registro data-first o narrativo).
Salida: Correos personalizados - difusion Inseguridad Mexico.docx (en esta misma carpeta).
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ---------- paleta ----------
AMBER  = RGBColor(0xB4, 0x7A, 0x10)
GREY   = RGBColor(0x55, 0x5B, 0x63)
INK    = RGBColor(0x1A, 0x1E, 0x24)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
ORANGE = RGBColor(0xB4, 0x60, 0x10)
SERIF = "Georgia"
SANS  = "Calibri"

URL_TABLERO   = "https://45digitalnoticias.github.io/Inseguridad-Mexico/index.html"
URL_COLUMNAS  = "https://45digitalnoticias.github.io/Inseguridad-Mexico/columnas/index.html"
URL_GOTHAM    = "https://45digitalnoticias.github.io/columnas/el-gotham-que-no-cuadra.html"

FIRMA = ["Sergio Rubén Valdespín Oseguera", "45 Digital Noticias", "facebook.com/45DigitalMx"]

SUBJ_DATA = "El −46% de homicidio doloso, desarmado con los datos del propio SESNSP (tablero abierto, 32 estados)"
SUBJ_NARR = "¿Cayó el homicidio 46% u 8.6%? Un tablero abierto lo desarma, estado por estado"

# ---------- palancas compartidas ----------
PALANCAS = [
 ("La curva.", "Una meseta de seis años (29,091 carpetas en 2018; 25,463 en 2024, incluso +0.6% ese año) que de pronto se desploma 21.5% en 2025. Sin pendiente previa, no describe una pacificación: describe un cambio en la forma de contar."),
 ("Las etiquetas.", "El estreno del Registro Nacional de Incidencia Delictiva (RNID) en febrero de 2026: robo de vehículo (40,866), narcomenudeo (35,074) y extorsión (3,683) pasan a cero en la serie histórica al fragmentarse en subtipos nuevos."),
 ("El promedio.", "El −21.5% nacional entierra a las cinco entidades que subieron, encabezadas por Sinaloa (+68%, de 492 a 1,355 carpetas), su año más violento desde 2011."),
 ("La fuga.", "2025 fue el peor año del registro de personas desaparecidas: el cuerpo que no se cuenta como muerto se cuenta, si acaso, como ausente, en un registro que no entra en el tablero de seguridad."),
]
SINTESIS = ("Sumadas, la misma baja admite tres lecturas: 46% (gobierno), 21.5% (carpetas a lo largo del año) y "
            "8.6% (México Evalúa, con la década en +68.2%). La diferencia no está en las calles: está en la elección "
            "de la ventana, el cambio de etiquetas y la aritmética del promedio.")
USO_LIBRE = ("El tablero y sus datos son de uso libre: cada cifra es trazable y vive en un archivo público (cifras.csv). "
             "Pueden citarla, enlazarla o reutilizarla con libertad; solo les pido acreditar la fuente, 45 Digital Noticias.")

def columnas_par(trato):
    enc = "encontrarán" if trato == "pl" else "encontrará"
    return (f"En el apartado de Columnas y auditorías {enc} la lectura firmada y la pinza «el anuncio vs. el archivo»: "
            "algunos comunicados oficiales contrastados contra SESNSP, INEGI y el padrón de desaparecidos. "
            "La tesis cabe en una sola lectura, «El Gotham que no cuadra».")

def intro_data(trato):
    les = "Les" if trato == "pl" else "Le"
    return (f"{les} escribo de 45 Digital Noticias. Publicamos un tablero público y gratuito que mapea la inseguridad "
            "en los 32 estados con los datos abiertos del SESNSP, y que desarma una cifra conocida: el −46% de homicidio "
            "doloso que presume el gobierno federal. Es verdadera en su aritmética y engañosa en lo que esconde.")

def intro_narr(trato):
    les = "Les" if trato == "pl" else "Le"
    return (f"{les} escribo de 45 Digital Noticias. A García Harfuch la prensa le puso un apodo que conviene tomar en "
            "serio: Batman. Y la metáfora rinde, porque Gotham ya ardía antes de que llegara el héroe. El gobierno presume "
            "que el homicidio doloso cayó 46%, su cifra más alta. Es verdadera en su aritmética y engañosa en lo que esconde. "
            "Publicamos un tablero público y gratuito que la desarma con los datos del propio SESNSP, estado por estado.")

PUENTE = "El tablero permite revisar, estado por estado y mes a mes (serie 2015 a 2026), las cuatro palancas que sostienen ese número:"

# ---------- DIRECTORIO con personalizacion ----------
# cat, nombre, correo, conf(OK/REV), trato(pl/sg), arch(data/narr), porque, subj(opcional)
C = [
 # ---------------- SOCIEDAD CIVIL Y OBSERVATORIOS ----------------
 ("Sociedad civil y observatorios",
  "Observatorio Nacional Ciudadano (ONC)", "contacto@onc.org.mx", "OK", "pl", "data",
  "Lo comparto con ustedes porque el ONC es el interlocutor natural de un ejercicio así: monitorea la incidencia con "
  "metodología propia y coordina la Red Nacional de Observatorios. Me interesa su lectura crítica; si detectan un error de "
  "interpretación o un dato mal trazado, quiero corregirlo.", None),

 ("Sociedad civil y observatorios",
  "México Evalúa", "info@mexicoevalua.org", "REV", "pl", "data",
  "Lo comparto con ustedes con una razón concreta: su estimación de −8.6% en violencia letal (con la década en +68.2%) es "
  "uno de los tres relojes con los que el tablero contrasta el −46% oficial. Les presento el ejercicio completo, apoyado en "
  "su lectura, aterrizado estado por estado.", "Su −8.6% como uno de los tres relojes del homicidio: tablero abierto, 32 estados"),

 ("Sociedad civil y observatorios",
  "México Unido Contra la Delincuencia (MUCD)", "comunicacion@mucd.org.mx", "OK", "pl", "data",
  "Lo comparto con ustedes porque MUCD ha insistido por años en que la política de seguridad se mide con datos, no con "
  "conferencias. El tablero pone esos datos al alcance de cualquiera, estado por estado, y abiertos a su escrutinio.", None),

 ("Sociedad civil y observatorios",
  "Impunidad Cero", "contacto@impunidadcero.org", "OK", "pl", "data",
  "Lo comparto con ustedes porque Impunidad Cero ha documentado cómo lo que no se cuenta se vuelve impune. La cuarta palanca "
  "del tablero, la fuga hacia el registro de desaparecidos, es exactamente ese problema medido.", None),

 ("Sociedad civil y observatorios",
  "World Justice Project México", "press@worldjusticeproject.mx", "OK", "pl", "data",
  "Lo comparto con ustedes porque el trabajo de WJP sobre Estado de Derecho e impunidad dialoga directo con la pregunta del "
  "tablero: qué tan confiable es la cifra con la que el propio Estado se evalúa.", None),

 ("Sociedad civil y observatorios",
  "Intersecta", "contacto@intersecta.org", "OK", "pl", "data",
  "Lo comparto con ustedes porque Intersecta analiza la violencia con los microdatos del SESNSP e INEGI. El tablero documenta "
  "cómo el cambio de metodología (RNID) rompe la comparabilidad de esas mismas series.", None),

 ("Sociedad civil y observatorios",
  "Data Cívica", "info@datacivica.org", "REV", "pl", "data",
  "Lo comparto con ustedes porque Data Cívica trabaja datos de violencia y desaparición con metodología abierta. El tablero "
  "también publica su CSV trazable, y la cuarta palanca (desaparecidos) es terreno que ustedes conocen mejor que nadie.", None),

 ("Sociedad civil y observatorios",
  "Causa en Común", "donaciones@causaencomun.org.mx", "REV", "pl", "data",
  "Lo comparto con ustedes porque Causa en Común ha sostenido el registro de atrocidades cuando el Estado dejó de hacerlo. "
  "El tablero comparte ese espíritu: que la cifra no entierre lo que arde.", None),

 ("Sociedad civil y observatorios",
  "Elementa DDHH", "info.mexico@elementaddhh.org", "OK", "pl", "data",
  "Lo comparto con ustedes porque Elementa trabaja la política de drogas y su violencia asociada. El tablero muestra cómo el "
  "narcomenudeo desaparece de la serie histórica (de 35,074 carpetas a cero) por reclasificación, no por reducción real.", None),

 ("Sociedad civil y observatorios",
  "Reinserta", "contacto@reinserta.org", "OK", "pl", "data",
  "Lo comparto con ustedes porque Reinserta conoce el costo humano detrás de cada cifra. Este tablero busca que esa cifra no "
  "se maquille en el camino.", None),

 ("Sociedad civil y observatorios",
  "TResearch (Carlos Penna)", "carlos@tresearch.mx", "OK", "sg", "data",
  "Lo comparto con usted porque TResearch levanta percepción y datos sobre homicidios y seguridad. El tablero ofrece la serie "
  "dura, estado por estado, para contrastar percepción con registro.", None),

 ("Sociedad civil y observatorios",
  "CIDAC", "info@cidac.org", "REV", "pl", "data",
  "Lo comparto con ustedes porque el proyecto de justicia de CIDAC mide procuración e impartición. El tablero aporta el insumo "
  "previo, la incidencia, y el problema de su comparabilidad tras el RNID.", None),

 ("Sociedad civil y observatorios",
  "Ethos Innovación en Políticas Públicas", "info@ethos.org.mx", "REV", "pl", "data",
  "Lo comparto con ustedes porque Ethos analiza políticas públicas con evidencia, incluida la economía del crimen. El tablero "
  "ofrece la serie de incidencia, abierta y trazable, para alimentar ese análisis.", None),

 ("Sociedad civil y observatorios",
  "Lantia Intelligence", "contacto@lantiaintelligence.com", "REV", "pl", "data",
  "Lo comparto con ustedes porque Lantia rastrea la violencia y el crimen organizado con fuentes abiertas. El caso Sinaloa del "
  "tablero (de 492 a 1,355 carpetas, +68%) dialoga con su seguimiento de la guerra entre los Chapitos y los Mayitos.", None),

 # ---------------- ACADEMIA Y THINK TANKS ----------------
 ("Academia y think tanks",
  "INSYDE — Instituto para la Seguridad y la Democracia", "info@insyde.org.mx", "OK", "pl", "data",
  "Lo comparto con ustedes porque INSYDE ha empujado la reforma policial sobre evidencia. El tablero documenta cómo la métrica "
  "con la que se evalúa esa política cambió de reglas a mitad del partido.", None),

 ("Academia y think tanks",
  "ITAM — CESIG", "cesig@itam.mx", "OK", "pl", "data",
  "Lo comparto con ustedes porque el CESIG forma analistas de seguridad e inteligencia. El tablero puede servir como caso de "
  "estudio sobre cómo una serie estadística se vuelve simulacro.", None),

 ("Academia y think tanks",
  "Ibero — Programa de Seguridad Ciudadana", "psc@ibero.mx", "OK", "pl", "data",
  "Lo comparto con ustedes porque el Programa de Seguridad Ciudadana trabaja la «vía civil». El tablero ofrece a esa mirada un "
  "insumo abierto y trazable sobre la incidencia real.", None),

 ("Academia y think tanks",
  "CIDE — División de Estudios Jurídicos", "ximena.medellin@cide.edu", "OK", "pl", "data",
  "Lo comparto con ustedes porque la División de Estudios Jurídicos del CIDE investiga el sistema penal y la política de drogas. "
  "El tablero muestra cómo delitos enteros (narcomenudeo, extorsión) salen de la serie por reclasificación.", None),

 ("Academia y think tanks",
  "CIDE — Programa de Política de Drogas", "programadepoliticadedrogas@gmail.com", "REV", "pl", "data",
  "Lo comparto con ustedes porque el Programa de Política de Drogas conoce el peso del narcomenudeo en la estadística. El tablero "
  "documenta su desaparición de la serie histórica (de 35,074 carpetas a cero) por cambio de etiqueta.", None),

 ("Academia y think tanks",
  "El Colef (difusión)", "difusionuec@colef.mx", "OK", "pl", "data",
  "Lo comparto con ustedes porque El Colef estudia la violencia y la seguridad en la frontera. El tablero permite ver el "
  "contraste estado por estado, con especial atención a entidades como Sinaloa.", None),

 ("Academia y think tanks",
  "COLMEX — Centro de Estudios Sociológicos", "direccion.ces@colmex.mx", "OK", "pl", "data",
  "Lo comparto con ustedes porque el CES estudia la violencia y el conflicto social desde la sociología. El tablero aporta la "
  "serie cuantitativa abierta para ese análisis.", None),

 ("Academia y think tanks",
  "Instituto Belisario Domínguez (Senado)", "apoyotecnico.ibd@senado.gob.mx", "OK", "pl", "data",
  "Lo comparto con ustedes porque el IBD produce investigación estratégica para el Senado. El tablero ofrece, con datos del "
  "propio SESNSP, una lectura de por qué tres metodologías arrojan tres cifras distintas (46, 21.5 y 8.6 por ciento).", None),

 ("Academia y think tanks",
  "Integralia Consultores", "contacto@integralia.com.mx", "OK", "pl", "data",
  "Lo comparto con ustedes porque Integralia evalúa el riesgo político y social. El tablero aporta una lectura crítica de la "
  "cifra de seguridad con la que el gobierno construye su narrativa.", None),

 ("Academia y think tanks",
  "UNAM — Instituto de Investigaciones Sociales", "iis@sociales.unam.mx", "OK", "pl", "data",
  "Lo comparto con ustedes porque el IIS investiga la violencia y el conflicto social. El tablero pone a disposición la serie "
  "2015 a 2026 con su CSV trazable.", None),

 ("Academia y think tanks",
  "UNAM — CISAN", "cisan@unam.mx", "OK", "pl", "data",
  "Lo comparto con ustedes porque el CISAN analiza la seguridad regional y la relación con Estados Unidos. El caso Sinaloa del "
  "tablero, tras la captura de «El Mayo» Zambada, conecta con esa agenda.", None),

 # ---------------- MEDIOS NACIONALES ----------------
 ("Medios nacionales",
  "Carlos Loret de Mola (El Universal / Latinus)", "historiasreportero@gmail.com", "OK", "sg", "narr",
  "Se lo comparto a usted en particular porque ha trabajado de cerca las cifras de Sinaloa, donde el homicidio diario pasó de "
  "1.4 a 6.9. El tablero ofrece esa serie, estado por estado, y la lectura de por qué el −46% nacional no cuadra con lo que arde "
  "en Culiacán. Por si le sirve para una columna o para Latinus.",
  "El −46% que no cuadra con Culiacán: un tablero abierto desarma la cifra, estado por estado"),

 ("Medios nacionales",
  "El Economista (redacción)", "redaccion@eleconomista.mx", "OK", "pl", "narr",
  "Lo comparto con su redacción por si resulta material para una nota o columna de análisis de datos: todo el tablero es abierto "
  "y trazable, listo para verificar cualquier cifra.", None),

 ("Medios nacionales",
  "Aristegui Noticias (redacción)", "redaccion@aristeguinoticias.com", "OK", "pl", "narr",
  "Lo comparto con su redacción por si les resulta de interés periodístico: el ejercicio desarma, con los propios datos del "
  "Estado, la cifra que sostiene la narrativa de seguridad del gobierno.", None),

 ("Medios nacionales",
  "El Universal (investigación / seguridad)", "denuncia@eluniversal.com.mx", "OK", "pl", "narr",
  "Lo comparto con el área de investigación por si abre una línea: el tablero documenta, dato por dato, cómo se construye el "
  "−46% y qué esconde el promedio nacional.", None),

 ("Medios nacionales",
  "El Universal (colaboraciones / contacto)", "contacto@eluniversal.com.mx", "OK", "pl", "narr",
  "Lo comparto con ustedes por si resulta material para sus páginas. La lectura editorial completa, «El Gotham que no cuadra», "
  "está enlazada abajo.", None),

 ("Medios nacionales",
  "Quinto Elemento Lab", "quintoelab@gmail.com", "OK", "pl", "data",
  "Lo comparto con ustedes porque su trabajo de investigación con datos encaja con el ejercicio. Queda abierto y trazable por "
  "si les sirve de insumo para una historia.", None),

 ("Medios nacionales",
  "Infobae México (redacción)", "redaccion@infobae.com", "OK", "pl", "narr",
  "Lo comparto con su redacción por si resulta de interés: una lectura, con datos del propio SESNSP, de por qué la baja del "
  "homicidio admite tres cifras muy distintas.", None),

 ("Medios nacionales",
  "Expansión Política (edición)", "mibarra@grupoexpansion.com", "OK", "sg", "narr",
  "Lo comparto con la sección de Política por si resulta material de análisis: el tablero desarma la cifra de seguridad con la "
  "que el gobierno construye su narrativa, estado por estado.", None),

 ("Medios nacionales",
  "Eje Central (redacción)", "comenta@ejecentral.com.mx", "OK", "pl", "narr",
  "Lo comparto con ustedes por si les resulta de interés: la cifra del −46% desarmada con los propios datos abiertos del "
  "SESNSP.", None),

 ("Medios nacionales",
  "Latinus (colaboraciones)", "contacto@latinus.us", "OK", "pl", "narr",
  "Lo comparto con ustedes por si resulta material para sus espacios: un tablero abierto que contrasta el dato oficial contra "
  "el archivo del propio Estado.", None),

 ("Medios nacionales",
  "Excélsior (comentarios)", "mvelazquezm@gimm.com.mx", "OK", "pl", "narr",
  "Lo comparto con ustedes por si resulta de interés periodístico: la lectura completa, «El Gotham que no cuadra», está "
  "enlazada abajo, con todo el tablero abierto detrás.", None),

 ("Medios nacionales",
  "Milenio (digital)", "digital@milenio.com", "OK", "pl", "narr",
  "Lo comparto con su redacción digital por si resulta material: un tablero interactivo, abierto y trazable, sobre la "
  "inseguridad en los 32 estados.", None),

 ("Medios nacionales",
  "Animal Político (El Sabueso)", "elsabueso@animalpolitico.com", "REV", "pl", "data",
  "Lo comparto con El Sabueso con una razón directa: verifican afirmaciones del poder, y el −46% es justo el tipo de cifra "
  "verificable. El tablero la desarma con los datos del propio SESNSP, palanca por palanca.", None),

 ("Medios nacionales",
  "Reforma (cartas / opinión)", "cartas@reforma.com", "REV", "pl", "narr",
  "Lo comparto con ustedes por si resulta material para sus páginas de opinión: una lectura, con datos abiertos del SESNSP, de "
  "por qué la baja del homicidio no cuadra.", None),

 ("Medios nacionales",
  "El Financiero (contacto web)", "contactoweb@elfinanciero.com.mx", "REV", "pl", "narr",
  "Lo comparto con su redacción por si resulta de interés: el tablero traduce a datos, estado por estado, la cifra de seguridad "
  "que sostiene la narrativa oficial.", None),

 # ---------------- MORELOS Y REGIONAL ----------------
 ("Morelos y regional",
  "Observatorio de Seguridad Ciudadana (UAEM)", "obs-seguridad@obs-seguridad.org", "OK", "pl", "data",
  "Lo comparto con ustedes porque son el interlocutor natural en Morelos: el tablero incluye el caso del estado a nivel municipal "
  "y audita el −62.8% que anunció García Harfuch (27 de mayo) contra los archivos del SESNSP, INEGI y el RNPDNO. Me interesa su "
  "lectura crítica.", "Morelos a nivel municipal y el −62.8% auditado: tablero abierto de inseguridad"),

 ("Morelos y regional",
  "IMIPE Morelos — Unidad de Transparencia", "imipe.ut@gmail.com", "OK", "pl", "data",
  "Lo comparto con ustedes porque el tablero descansa en datos públicos y trazables, en el espíritu del derecho de acceso a la "
  "información. Incluye un caso de Morelos a nivel municipal y la auditoría de comunicados oficiales contra los registros del "
  "propio Estado.", "Transparencia hecha tablero: inseguridad en Morelos y a nivel nacional, con datos abiertos"),

 ("Morelos y regional",
  "IMIPE Morelos — Socialización / Capacitación", "imipe.socializacion@gmail.com", "OK", "pl", "data",
  "Lo comparto con ustedes porque el tablero es una herramienta ciudadana de lectura de datos abiertos que puede ser útil para "
  "sus labores de socialización y capacitación. Es de uso libre y todo dato es trazable.",
  "Herramienta ciudadana de datos abiertos: tablero de inseguridad (uso libre)"),

 ("Morelos y regional",
  "El Sol de Cuernavaca", "tuopinion@elsoldecuernavaca.com.mx", "OK", "pl", "narr",
  "Lo comparto con su redacción porque a un medio de Morelos esto le pega de cerca: el tablero incluye el estado a nivel "
  "municipal y audita el −62.8% que anunció García Harfuch contra los archivos del SESNSP, INEGI y el RNPDNO.",
  "Morelos y el −62.8%: un tablero abierto audita la cifra contra el archivo del Estado"),

 ("Morelos y regional",
  "El Regional del Sur", "contacto@elregional.com.mx", "OK", "pl", "narr",
  "Lo comparto con su redacción porque a un medio de Morelos esto le toca de cerca: el tablero trae el caso del estado a nivel "
  "municipal y la auditoría del −62.8% anunciado para Morelos contra los registros oficiales.",
  "Morelos en el tablero: el −62.8% auditado contra el archivo del Estado"),

 ("Morelos y regional",
  "Diario de Morelos (redacción)", "redaccion@diariodemorelos.com", "REV", "pl", "narr",
  "Lo comparto con su redacción porque a un medio estatal esto le pega de cerca: el tablero incluye Morelos a nivel municipal y "
  "audita el −62.8% que se anunció para el estado contra los archivos del SESNSP, INEGI y el RNPDNO.",
  "Morelos y el −62.8%: tablero abierto que audita la cifra contra el archivo"),

 ("Morelos y regional",
  "La Unión de Morelos (redacción)", "uniondemorelos@gmail.com", "REV", "pl", "narr",
  "Lo comparto con su redacción porque a un medio de Morelos esto le toca de cerca: el tablero trae el estado a nivel municipal "
  "y la auditoría del −62.8% anunciado para Morelos contra los registros oficiales.",
  "Morelos en el tablero: el −62.8% auditado contra el archivo del Estado"),

 ("Morelos y regional",
  "La Jornada Morelos (redacción)", "redaccionjornadamorelos@gmail.com", "REV", "pl", "narr",
  "Lo comparto con su redacción porque a un medio de Morelos esto le pega de cerca: el tablero incluye el estado a nivel "
  "municipal y audita el −62.8% anunciado para Morelos contra los archivos del SESNSP, INEGI y el RNPDNO.",
  "Morelos y el −62.8%: tablero abierto que audita la cifra contra el archivo"),

 ("Morelos y regional",
  "El Sol de Cuautla (redacción)", "elsoldecuautla@yahoo.com.mx", "REV", "pl", "narr",
  "Lo comparto con su redacción porque a un medio de Morelos esto le toca de cerca: el tablero trae el caso del estado a nivel "
  "municipal y la auditoría del −62.8% anunciado para Morelos.",
  "Morelos en el tablero: el −62.8% auditado contra el archivo del Estado"),

 ("Morelos y regional",
  "Ana María Salazar (periodista de seguridad)", "anamariasalazar@gmail.com", "REV", "sg", "narr",
  "Se lo comparto a usted porque su análisis de seguridad y justicia es referencia obligada. El tablero ofrece la serie dura, "
  "estado por estado, y la lectura de por qué el −46% admite tres cifras muy distintas.", None),
]

# Contactos SIN correo (solo se listan)
SIN_CORREO = [
 ("Semáforo Delictivo", "Solo app / Facebook oficial"),
 ("Jorge Medellín (EstadoMayor.mx)", "Solo X: @EstadoMayor_mx"),
 ("Zona Centro Noticias", "Solo redes sociales"),
]

# ================= construccion del documento =================
doc = Document()
for s in doc.sections:
    s.top_margin = Inches(0.8); s.bottom_margin = Inches(0.8)
    s.left_margin = Inches(0.9); s.right_margin = Inches(0.9)
base = doc.styles["Normal"]; base.font.name = SANS; base.font.size = Pt(10.5)

def add_runs(p, segments):
    for text, bold, color, size in segments:
        r = p.add_run(text); r.bold = bold
        r.font.color.rgb = color; r.font.size = Pt(size); r.font.name = SANS

def para(segments, justify=True, space_after=8, space_before=0, indent=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if indent: p.paragraph_format.left_indent = Inches(indent)
    add_runs(p, segments)
    return p

def hrule(color="B47A10", sz="10"):
    fl = doc.add_paragraph(); pPr = fl._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), color)
    pbdr.append(bottom); pPr.append(pbdr)
    fl.paragraph_format.space_after = Pt(6)

def page_break():
    doc.add_page_text = None
    p = doc.add_paragraph(); run = p.add_run(); br = OxmlElement("w:br")
    br.set(qn("w:type"), "page"); run._r.append(br)

# ---------- portada ----------
para([("45 DIGITAL NOTICIAS", True, AMBER, 10)], justify=False, space_after=2)
para([("Correos de difusión personalizados", True, INK, 22)], justify=False, space_after=2)
para([("Dashboard · Inseguridad México · 32 estados", True, AMBER, 12)], justify=False, space_after=6)
hrule()
para([("Un correo por objetivo, listo para copiar y pegar · 16 de junio de 2026", False, GREY, 9.5)], justify=False, space_after=10)
para([("Cuerpo de datos compartido (las cuatro palancas + las tres lecturas), personalizado en saludo, "
       "asunto y el párrafo de «por qué usted/ustedes». El tablero se declara de uso libre con atribución a la fuente.",
       False, GREY, 9.5)], space_after=6)
para([("Antes de enviar: confirme los correos marcados «verificar» (REV). Personalice el saludo si tiene un "
       "nombre de pila. Adjunte, si quiere, una captura del gráfico de la meseta.", False, ORANGE, 9.5)], space_after=4)

total = len(C)
cur_cat = None
idx = 0
for cat, nombre, correo, conf, trato, arch, porque, subj in C:
    idx += 1
    page_break()
    # encabezado de la ficha
    para([(cat.upper(), True, AMBER, 9)], justify=False, space_after=2)
    flag = "  ·  VERIFICAR CORREO ANTES DE ENVIAR" if conf == "REV" else ""
    para([("Para: ", True, INK, 10.5), (nombre, False, INK, 10.5)], justify=False, space_after=1)
    para([("Correo: ", True, INK, 10.5), (correo, False, AMBER, 10.5),
          (flag, True, ORANGE, 8.5)], justify=False, space_after=1)
    s = subj if subj else (SUBJ_NARR if arch == "narr" else SUBJ_DATA)
    para([("Asunto: ", True, INK, 10.5), (s, False, INK, 10.5)], justify=False, space_after=6)
    hrule(color="D0D4DA", sz="6")

    # saludo
    if trato == "sg":
        saludo = f"Estimado/a {nombre.split(' (')[0]}:"
    else:
        saludo = f"Estimadas y estimados de {nombre.split(' (')[0].split(' —')[0]}:"
    para([(saludo, False, INK, 10.5)], space_after=8)

    # intro
    intro = intro_narr(trato) if arch == "narr" else intro_data(trato)
    para([(intro, False, INK, 10.5)], space_after=8)

    # puente
    para([(PUENTE, False, INK, 10.5)], space_after=6)

    # palancas
    for i, (lead, txt) in enumerate(PALANCAS, 1):
        para([(f"{i}. ", True, AMBER, 10.5), (lead + " ", True, INK, 10.5), (txt, False, INK, 10.5)],
             space_after=5, indent=0.18)

    # sintesis
    para([(SINTESIS, False, INK, 10.5)], space_after=8, space_before=2)
    # columnas
    para([(columnas_par(trato), False, INK, 10.5)], space_after=8)
    # porque (personalizado)
    para([(porque, False, INK, 10.5)], space_after=8)
    # uso libre
    para([(USO_LIBRE, False, INK, 10.5)], space_after=8)

    # enlaces
    para([("Tablero: ", True, INK, 9.5), (URL_TABLERO, False, AMBER, 9.5)], justify=False, space_after=1)
    para([("Columnas y auditorías: ", True, INK, 9.5), (URL_COLUMNAS, False, AMBER, 9.5)], justify=False, space_after=1)
    para([("Columna «El Gotham que no cuadra»: ", True, INK, 9.5), (URL_GOTHAM, False, AMBER, 9.5)], justify=False, space_after=8)

    # cierre + firma
    para([("Gracias por su tiempo.", False, INK, 10.5)], space_after=8)
    para([(FIRMA[0], True, INK, 10.5)], justify=False, space_after=0)
    para([(FIRMA[1], False, GREY, 9.5)], justify=False, space_after=0)
    para([(FIRMA[2], False, GREY, 9.5)], justify=False, space_after=0)

# ---------- pagina final: sin correo ----------
page_break()
para([("CONTACTOS SIN CORREO PÚBLICO (otra vía)", True, AMBER, 10)], justify=False, space_after=6)
for nombre, via in SIN_CORREO:
    para([("• ", True, AMBER, 10.5), (nombre + ": ", True, INK, 10.5), (via, False, GREY, 10.5)],
         justify=False, space_after=3)
para([(f"Total de correos personalizados en este documento: {total}.", False, GREY, 9.5)],
     justify=False, space_before=10)

out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Correos personalizados - difusion Inseguridad Mexico.docx")
doc.save(out)
print("OK ->", out, "|", total, "correos")
